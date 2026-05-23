from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "episode5_knowledge_mining_lab.docx"

LETTER_WIDTH_DXA = 12240
MARGIN_DXA = 1440
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120

BLUE = RGBColor(46, 116, 181)
DEEP_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(11, 37, 69)
MUTED = RGBColor(89, 89, 89)
TABLE_FILL = "E8EEF5"
RULE = "D7E3F2"


def set_font(run, size, color=INK, bold=False, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_indent(table, indent_dxa: int) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")


def set_table_layout(table) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in [("top", top), ("bottom", bottom), ("start", start), ("end", end)]:
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def apply_page_setup(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    pf = normal.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DEEP_BLUE, 10, 5),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15


def add_footer(section) -> None:
    paragraph = section.footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Episode 5 Knowledge Mining Lab")
    set_font(run, 9, color=MUTED, bold=False)


def add_title_block(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("EPISODE 5 KNOWLEDGE MINING LAB")
    set_font(run, 22, color=INK, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(14)
    run = subtitle.add_run(
        "Azure AI Search, Document Intelligence, and Content Understanding for the High-End Restaurant AI Automation solution"
    )
    set_font(run, 11, color=MUTED, bold=False)

    meta = doc.add_paragraph()
    meta.paragraph_format.space_before = Pt(0)
    meta.paragraph_format.space_after = Pt(12)
    meta_run = meta.add_run(
        "Audience: AI-102 practice builders | Scope: Episode 5 retrieval, ingestion, extraction, and analyzer workflows"
    )
    set_font(meta_run, 10.5, color=DEEP_BLUE, bold=False)

    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(0)
    rule.paragraph_format.space_after = Pt(12)
    p_pr = rule._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), RULE)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_intro(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.add_run(
        "This lab package extends the restaurant solution with a focused knowledge-mining layer. "
        "It gives you a practical path for testing retrieval quality, document extraction, and analyzer-driven multimodal review "
        "without having to redesign the earlier agentic or multilingual application slices."
    )


def add_endpoint_table(doc: Document) -> None:
    doc.add_paragraph("What Is Included", style="Heading 1")
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_layout(table)
    set_table_indent(table, TABLE_INDENT_DXA)
    headers = [("Capability", 2520), ("Episode 5 Surface", 6840)]
    for index, (label, width) in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_width(cell, width)
        set_cell_margins(cell)
        shade_cell(cell, TABLE_FILL)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        para = cell.paragraphs[0]
        para.paragraph_format.space_after = Pt(0)
        run = para.add_run(label)
        set_font(run, 10.5, color=INK, bold=True)

    rows = [
        ("Search", "POST /search/query with keyword, vector, hybrid, and semantic-style retrieval"),
        ("Operations", "GET /search/ingest-status for a quick readiness and corpus check"),
        ("Receipts", "POST /document/receipt for line items, totals, and merchant metadata"),
        ("Layout", "POST /document/layout for paragraphs, tables, and page-level structure"),
        ("Contracts", "POST /document/contract for private event extraction and review routing"),
        ("Analyzers", "POST /content-understanding/analyze for menu, contract, audio, video, and invoice analyzer patterns"),
    ]
    for left, right in rows:
        row = table.add_row()
        for index, (text, width) in enumerate([(left, 2520), (right, 6840)]):
            cell = row.cells[index]
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            para = cell.paragraphs[0]
            para.paragraph_format.space_after = Pt(0)
            run = para.add_run(text)
            set_font(run, 10.5, color=INK, bold=index == 0)


def add_numbered_steps(doc: Document) -> None:
    doc.add_paragraph("Suggested Lab Flow", style="Heading 1")
    steps = [
        "Keep MOCK_MODE=true and validate the new endpoints locally before connecting live Azure services.",
        "Run scripts/build_search_index.py to write the richer index schema and, when configured, create the Azure AI Search index.",
        "Run scripts/ingest_documents.py to chunk restaurant source files, generate embeddings, and upload enriched search records.",
        "Validate retrieval with questions about private dining minimums, vegan tasting availability, and shellfish references.",
        "Upload a receipt, invoice, or event contract to confirm extraction quality and review-threshold behavior.",
        "Test the content-understanding analyzers with representative PDF, audio, and video inputs.",
    ]
    for step in steps:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.25
        paragraph.add_run(step)


def add_defaults_table(doc: Document) -> None:
    doc.add_paragraph("Retrieval Defaults", style="Heading 1")
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_layout(table)
    set_table_indent(table, TABLE_INDENT_DXA)
    widths = [2520, 1800, 5040]
    labels = ["Setting", "Default", "Usage Note"]
    for index, (label, width) in enumerate(zip(labels, widths)):
        cell = table.rows[0].cells[index]
        set_cell_width(cell, width)
        set_cell_margins(cell)
        shade_cell(cell, TABLE_FILL)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        run = cell.paragraphs[0].add_run(label)
        set_font(run, 10.5, color=INK, bold=True)

    entries = [
        ("Query mode", "hybrid", "Balances lexical matching with vector-style grounding."),
        ("Result count", "5", "Keeps answers concise while still supporting citations."),
        ("Semantic config", "restaurant-semantic", "Used when semantic or hybrid ranking is enabled."),
        ("Vector size", "3072", "Matches the embedding shape used by the search ingestion flow."),
        ("Common filters", "menu / policy / pairings / vegan", "Use document_type, menu_section, and allergen_tag to narrow recall."),
    ]
    for row_values in entries:
        row = table.add_row()
        for index, (text, width) in enumerate(zip(row_values, widths)):
            cell = row.cells[index]
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            para = cell.paragraphs[0]
            para.paragraph_format.space_after = Pt(0)
            run = para.add_run(text)
            set_font(run, 10.5, color=INK, bold=index == 0)


def add_why_it_matters(doc: Document) -> None:
    doc.add_paragraph("Why This Matters for AI-102", style="Heading 1")
    points = [
        "Azure AI Search index design, chunking strategy, and citation-friendly retrieval outputs.",
        "Grounding prompts with explicit sources instead of relying on unsupported model memory.",
        "Document Intelligence extraction for structured and semi-structured restaurant operations documents.",
        "Separation of deterministic extraction logic from higher-level generative orchestration.",
        "Operational concerns such as ingest status, review thresholds, and metadata-driven filtering.",
    ]
    for point in points:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.25
        paragraph.add_run(point)


def build_doc() -> None:
    doc = Document()
    apply_page_setup(doc)
    configure_styles(doc)
    add_footer(doc.sections[0])
    add_title_block(doc)
    add_intro(doc)
    add_endpoint_table(doc)
    add_numbered_steps(doc)
    add_defaults_table(doc)
    add_why_it_matters(doc)
    doc.save(OUTPUT)


if __name__ == "__main__":
    build_doc()
    print(OUTPUT)
